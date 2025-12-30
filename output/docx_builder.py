"""
Word document generation for affidavit output.
"""
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Tuple
import logging
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pipeline.core import PipelineState, ErrorSeverity

logger = logging.getLogger(__name__)


class AffidavitDocxBuilder:
    """Builds Word documents from pipeline output."""

    def __init__(self):
        self.doc = None

    def build(self, state: PipelineState) -> Tuple[str, str]:
        """
        Build two separate Word documents from pipeline state.

        Args:
            state: Final pipeline state

        Returns:
            Tuple of (main_document_path, technical_report_path)

        Raises:
            Exception: If document generation fails
        """
        try:
            # Build main document (draft + evaluation summary)
            main_path = self._build_main_document(state)

            # Build technical report (extraction details + processing info)
            report_path = self._build_technical_report(state)

            logger.info(f"Main document saved to: {main_path}")
            logger.info(f"Technical report saved to: {report_path}")

            return main_path, report_path

        except Exception as e:
            logger.error(f"Failed to build documents: {str(e)}")
            raise

    def _build_main_document(self, state: PipelineState) -> str:
        """Build the main affidavit document (clean, for attorney)."""
        self.doc = Document()
        self._setup_document()

        # Add title
        self._add_title("AFFIDAVIT DRAFT")

        # Add main affidavit body
        self._add_affidavit_body(state)

        # Add evaluation summary on same page
        self._add_evaluation_summary(state)

        # Save main document
        output_path = self._generate_output_path(
            state.output_path,
            state.case_name,
            suffix="draft"
        )
        self.doc.save(output_path)

        return output_path

    def _build_technical_report(self, state: PipelineState) -> str:
        """Build technical report with extraction and processing details."""
        self.doc = Document()
        self._setup_document()

        # Add title
        self._add_title("PROCESSING REPORT")

        # Add extraction details
        self._add_extraction_report(state)

        # Add detailed evaluation
        self._add_page_break()
        self._add_detailed_evaluation(state)

        # Add processing metadata
        self._add_page_break()
        self._add_processing_metadata(state)

        # Save report
        output_path = self._generate_output_path(
            state.output_path,
            state.case_name,
            suffix="report"
        )
        self.doc.save(output_path)

        return output_path

    def _setup_document(self):
        """Configure document defaults."""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        # Set margins (1 inch all around)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def _add_title(self, title_text: str):
        """Add document title."""
        title = self.doc.add_heading(title_text, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add generation timestamp
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        p = self.doc.add_paragraph(f"Generated: {timestamp}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()  # Blank line

    def _add_affidavit_body(self, state: PipelineState):
        """Add the main affidavit body text."""
        self.doc.add_heading('AFFIDAVIT BODY', level=2)

        # Use final_text if available, otherwise draft_text
        body_text = state.final_text or state.draft_text or "[No text generated]"

        # Add paragraphs
        for paragraph_text in body_text.split('\n\n'):
            if paragraph_text.strip():
                self.doc.add_paragraph(paragraph_text.strip())

    def _add_page_break(self):
        """Add page break."""
        self.doc.add_page_break()

    def _add_extraction_report(self, state: PipelineState):
        """Add component extraction report."""
        self.doc.add_heading('APPENDIX A: Extraction Report', level=2)

        if not state.extracted_components:
            self.doc.add_paragraph("[No extraction data available]")
            return

        self.doc.add_paragraph(
            "The following components were extracted from the interview notes:"
        )
        self.doc.add_paragraph()  # Blank line

        # Format as readable text
        for component, content in state.extracted_components.items():
            p = self.doc.add_paragraph(style='List Bullet')
            p.add_run(f"{component}: ").bold = True

            if isinstance(content, dict) or isinstance(content, list):
                content_str = json.dumps(content, indent=2)
            else:
                content_str = str(content)

            if content_str == "MISSING" or not content_str.strip():
                run = p.add_run("[MISSING]")
                run.font.color.rgb = None  # Red would require extra import
            else:
                p.add_run(content_str)

    def _add_evaluation_summary(self, state: PipelineState):
        """Add brief evaluation summary (for main document)."""
        self.doc.add_paragraph()  # Blank line
        self.doc.add_heading('Evaluation Summary', level=2)

        if not state.evaluation_report:
            self.doc.add_paragraph("Evaluation data not available.")
            return

        eval_report = state.evaluation_report

        # Summary statement
        if eval_report.get('all_supported', False):
            self.doc.add_paragraph(
                "All statements in the draft are supported by the source material."
            )
        else:
            unsupported = eval_report.get('unsupported_statements', [])
            uncertain = eval_report.get('uncertain_statements', [])

            summary_parts = []
            if unsupported:
                summary_parts.append(f"{len(unsupported)} unsupported statement(s)")
            if uncertain:
                summary_parts.append(f"{len(uncertain)} uncertain statement(s)")

            if summary_parts:
                self.doc.add_paragraph(
                    f"Note: {' and '.join(summary_parts)} were identified during evaluation. "
                    "See technical report for details."
                )

        # Brief metadata
        self.doc.add_paragraph(f"Revision iterations: {state.iteration_count}")

    def _add_detailed_evaluation(self, state: PipelineState):
        """Add detailed evaluation report (for technical report)."""
        self.doc.add_heading('Detailed Evaluation Report', level=2)

        if not state.evaluation_report:
            self.doc.add_paragraph("No evaluation data available.")
            return

        eval_report = state.evaluation_report

        # Overall status
        if eval_report.get('all_supported', False):
            self.doc.add_paragraph(
                "Status: All statements supported by source material."
            )
        else:
            self.doc.add_paragraph(
                "Status: Some statements require attention."
            )

        self.doc.add_paragraph()

        # Unsupported statements
        unsupported = eval_report.get('unsupported_statements', [])
        if unsupported:
            self.doc.add_heading('Unsupported Statements', level=3)
            self.doc.add_paragraph(
                f"Found {len(unsupported)} statement(s) not supported by source material:"
            )
            for stmt in unsupported:
                self.doc.add_paragraph(f"• {stmt}", style='List Bullet')
            self.doc.add_paragraph()

        # Uncertain statements
        uncertain = eval_report.get('uncertain_statements', [])
        if uncertain:
            self.doc.add_heading('Uncertain Statements', level=3)
            self.doc.add_paragraph(
                f"Found {len(uncertain)} statement(s) that may need verification:"
            )
            for stmt in uncertain:
                self.doc.add_paragraph(f"• {stmt}", style='List Bullet')
            self.doc.add_paragraph()

        # Summary from evaluation
        if 'summary' in eval_report:
            self.doc.add_heading('Evaluator Summary', level=3)
            self.doc.add_paragraph(eval_report['summary'])

    def _add_processing_metadata(self, state: PipelineState):
        """Add processing metadata and error log (for technical report)."""
        self.doc.add_heading('Processing Metadata', level=2)

        # Iterations
        self.doc.add_paragraph(f"Revision iterations completed: {state.iteration_count}")
        self.doc.add_paragraph(f"Maximum iterations allowed: 3")
        self.doc.add_paragraph()

        # Errors
        if state.errors:
            self.doc.add_heading('Processing Errors', level=3)
            self.doc.add_paragraph(
                f"{len(state.errors)} error(s) occurred during processing:"
            )
            for error in state.errors:
                p = self.doc.add_paragraph(style='List Bullet')
                severity_str = f"[{error.severity.value.upper()}]"
                p.add_run(f"{severity_str} {error.step_name}: ").bold = True
                p.add_run(error.message)
        else:
            self.doc.add_paragraph("No errors occurred during processing.")

    def _sanitize_case_name(self, case_name: str) -> str:
        """
        Sanitize case name for use as directory name.

        Args:
            case_name: Raw case name from user

        Returns:
            Safe directory name
        """
        # Convert to lowercase, replace spaces with underscores
        sanitized = case_name.lower().strip()
        sanitized = re.sub(r'\s+', '_', sanitized)
        # Remove any non-alphanumeric characters except underscore and hyphen
        sanitized = re.sub(r'[^a-z0-9_-]', '', sanitized)
        # Limit length
        sanitized = sanitized[:50]
        return sanitized or "case"

    def _generate_output_path(self, base_path: str, case_name: str, suffix: str = "draft") -> str:
        """
        Generate output path in case-specific subdirectory.

        Args:
            base_path: Base output directory
            case_name: Name of the case (creates subdirectory)
            suffix: Filename suffix ("draft" or "report")

        Returns:
            Full path to output file
        """
        # Create case subdirectory
        base_dir = Path(base_path)
        safe_case_name = self._sanitize_case_name(case_name)
        case_dir = base_dir / safe_case_name

        # Create directory if it doesn't exist
        case_dir.mkdir(parents=True, exist_ok=True)

        # Simple filenames without timestamp
        filename = f"{suffix}.docx"
        return str(case_dir / filename)
